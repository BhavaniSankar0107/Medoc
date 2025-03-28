import os
import streamlit as st
import tempfile
import torchaudio
import whisper
from llama_cpp import Llama
from docx import Document
from huggingface_hub import hf_hub_download

# Hugging Face model details
REPO_ID = "bhavanisankar-45/mistral"  # Replace with your Hugging Face repo
MODEL_FILENAME = "mistral-7b-instruct-v0.2.Q4_K_M.gguf"

# Download model from Hugging Face
st.write("Fetching AI model from Hugging Face... This may take a while.")
MODEL_PATH = hf_hub_download(repo_id=REPO_ID, filename=MODEL_FILENAME)
st.write("Model loaded successfully!")

# Load LLaMA model
llm = Llama(model_path=MODEL_PATH, n_ctx=1024)

def transcribe_audio(audio_path):
    """Transcribe audio using Whisper."""
    model = whisper.load_model("base")
    try:
        waveform, sample_rate = torchaudio.load(audio_path)
        result = model.transcribe(audio_path)
        return result["text"]
    except Exception as e:
        st.error(f"Error processing audio: {e}")
        return "Error processing audio."

def generate_soap_note(conversation):
    """Generate a structured SOAP note using the AI model."""
    prompt = f"""
    Convert the following conversation into a structured SOAP note.

    Patient Details:
    Name: Extract if mentioned, else "Not Mentioned"
    Age: Extract if mentioned, else "Not Mentioned"
    Gender: Extract if mentioned, else "Not Mentioned"

    Subjective:
    Main complaint and symptoms
    Medical history

    Objective:
    Physical exam findings or "Not Available"
    Test results or "Not Available"

    Assessment:
    Possible diagnosis
    Probable disease

    Plan:
    Medications, tests, and follow-up

    Conversation:
    {conversation}

    SOAP Note:
    """
    output = llm(prompt=prompt, max_tokens=400, temperature=0.3, stop=["###"])
    return output["choices"][0]["text"].strip()

def save_soap_note_to_word(soap_note):
    """Save the SOAP note as a Word document and return the file path."""
    doc = Document()
    doc.add_heading("SOAP Note", level=1)

    sections = soap_note.split("\n\n")
    for section in sections:
        if ":" in section:
            lines = section.split("\n")
            doc.add_heading(lines[0].strip(), level=2)
            for line in lines[1:]:
                if ":" in line:
                    doc.add_paragraph(f"• {line.strip()}")
                else:
                    doc.add_paragraph(line.strip())

    temp_word_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_word_file.name)
    return temp_word_file.name

# Streamlit UI
st.title("AI-Based Medical Documentation Assistant")

option = st.radio("Choose an option:", ("Upload Audio",))
audio_file = None
transcription = None
soap_note = None

if option == "Upload Audio":
    uploaded_file = st.file_uploader("Upload an audio file", type=["mp3", "wav", "m4a"])
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.read())
            audio_file = temp_file.name

        st.audio(audio_file, format="audio/wav")

        if st.button("Analyze"):
            st.write("Transcribing...")
            transcription = transcribe_audio(audio_file)
            st.subheader("Transcription:")
            st.write(transcription)

            st.write("Generating SOAP Note...")
            soap_note = generate_soap_note(transcription)
            st.subheader("SOAP Note:")
            st.write(soap_note)

            soap_note_file = save_soap_note_to_word(soap_note)
            with open(soap_note_file, "rb") as file:
                st.download_button(
                    label="Download SOAP Note",
                    data=file,
                    file_name="SOAP_Note.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        os.remove(audio_file)
